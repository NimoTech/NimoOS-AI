"""Document-to-Markdown extractors. Called from attachments.upload.handle_upload
on a worker thread (asyncio.to_thread); functions here are blocking and CPU
bound and MUST NOT be awaited directly.
"""
from __future__ import annotations

import logging

log = logging.getLogger(__name__)


def _zipbomb_check(path: str, max_uncompressed_bytes: int) -> bool:
    """True if the ZIP at `path` decompresses within budget. False otherwise.
    A non-ZIP (or unreadable ZIP) also returns False — caller treats that as
    a zip_bomb decision rather than crashing; the subsequent parser call will
    typically report parse_error if invoked, but for ZIP-based formats we
    bail early."""
    import zipfile
    try:
        with zipfile.ZipFile(path) as z:
            total = sum(zi.file_size for zi in z.infolist())
    except (zipfile.BadZipFile, OSError):
        return False
    return total <= max_uncompressed_bytes


def extract_to_markdown(path: str, ext: str, *,
                        max_chars: int,
                        max_uncompressed_bytes: int) -> dict:
    """Dispatch a document path to the matching format extractor.

    Returns:
      success: {"ok": True, "markdown": str, "pages": int|None,
                "chars": int, "truncated": bool, "extractor": str}
      failure: {"ok": False, "error": str}
        error in {not_installed, unsupported, zip_bomb,
                  parse_error, encrypted, empty_scanned}
    """
    ext = ext.lower()
    if ext == "pdf":
        return _extract_pdf(path, max_chars=max_chars)
    if ext == "docx":
        return _extract_docx(path, max_chars=max_chars,
                             max_uncompressed_bytes=max_uncompressed_bytes)
    if ext in ("xlsx", "xlsm"):
        return _extract_xlsx(path, max_chars=max_chars,
                             max_uncompressed_bytes=max_uncompressed_bytes)
    if ext == "pptx":
        return _extract_pptx(path, max_chars=max_chars,
                             max_uncompressed_bytes=max_uncompressed_bytes)
    return {"ok": False, "error": "unsupported"}


# Per-format implementations are added in later tasks. Stubs return
# not_installed so the dispatcher is fully exercisable today.
def _extract_pdf(path: str, *, max_chars: int) -> dict:
    try:
        import pypdf
    except ImportError:
        return {"ok": False, "error": "not_installed"}

    try:
        reader = pypdf.PdfReader(path)

        if reader.is_encrypted:
            # decrypt('') returns 0 when the empty password fails, 1/2 on success.
            try:
                ok = reader.decrypt("")
            except Exception:
                ok = 0
            if not ok:
                return {"ok": False, "error": "encrypted"}

        pages = list(reader.pages)
        buf: list[str] = []
        total = 0
        truncated = False
        for page in pages:
            try:
                raw = page.extract_text() or ""
            except Exception:
                # Per-page failure shouldn't sink the whole document; emit a
                # marker so the model knows that page was lost.
                raw = "[unreadable page]"
            text = raw.strip()
            if not text:
                continue
            if total + len(text) >= max_chars:
                buf.append(text[: max_chars - total])
                total = max_chars
                truncated = True
                break
            buf.append(text)
            total += len(text)
            buf.append("\n\n---\n\n")
            total += 7

        if not buf:
            return {"ok": False, "error": "empty_scanned"}

        markdown = "".join(buf).rstrip()

        return {
            "ok": True,
            "markdown": markdown,
            "pages": len(pages),
            "chars": len(markdown),
            "truncated": truncated,
            "extractor": "pypdf",
        }
    except Exception:
        log.exception("pypdf extraction failed for %s", path)
        return {"ok": False, "error": "parse_error"}


def _extract_docx(path: str, *, max_chars: int,
                  max_uncompressed_bytes: int) -> dict:
    try:
        import docx  # python-docx
    except ImportError:
        return {"ok": False, "error": "not_installed"}

    if not _zipbomb_check(path, max_uncompressed_bytes):
        return {"ok": False, "error": "zip_bomb"}

    try:
        d = docx.Document(path)
        buf: list[str] = []
        total = 0
        truncated = False

        def emit(line: str) -> bool:
            """Append `line + \n`; return True if we hit the cap."""
            nonlocal total, truncated
            if total + len(line) + 1 >= max_chars:
                buf.append(line[: max_chars - total])
                total = max_chars
                truncated = True
                return True
            buf.append(line + "\n")
            total += len(line) + 1
            return False

        for para in d.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            style = (para.style.name or "") if para.style is not None else ""
            if style.startswith("Heading"):
                # "Heading 1" → "# ", "Heading 2" → "## ", etc.
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 1
                level = max(1, min(level, 6))
                if emit("#" * level + " " + text):
                    break
            else:
                if emit(text):
                    break

        if not truncated:
            for table in d.tables:
                if truncated:
                    break
                rows = []
                for row in table.rows:
                    cells = [(cell.text or "").replace("|", "\\|").replace("\n", " ")
                             for cell in row.cells]
                    rows.append("| " + " | ".join(cells) + " |")
                if not rows:
                    continue
                header = rows[0]
                divider = "| " + " | ".join(["---"] * header.count("|") if header.count("|") > 1 else ["---"]) + " |"
                if emit(""): break
                if emit(header): break
                if emit(divider): break
                for body_row in rows[1:]:
                    if emit(body_row):
                        break

        markdown = "".join(buf).rstrip()
        if not markdown.strip():
            return {"ok": False, "error": "empty_scanned"}

        return {
            "ok": True,
            "markdown": markdown,
            "pages": None,
            "chars": len(markdown),
            "truncated": truncated,
            "extractor": "python-docx",
        }
    except Exception:
        log.exception("python-docx extraction failed for %s", path)
        return {"ok": False, "error": "parse_error"}


def _extract_xlsx(path: str, *, max_chars: int,
                  max_uncompressed_bytes: int) -> dict:
    try:
        import openpyxl
    except ImportError:
        return {"ok": False, "error": "not_installed"}

    if not _zipbomb_check(path, max_uncompressed_bytes):
        return {"ok": False, "error": "zip_bomb"}

    def _emit_sheet(ws, buf, total, truncated):
        """Append one sheet's markdown to buf. Returns
        (new_total, truncated, saw_any_value)."""
        saw_any_value = False
        rows_iter = ws.iter_rows(values_only=True)
        first_row = None
        for row in rows_iter:
            first_row = row
            break
        if first_row is None:
            return total, truncated, saw_any_value

        def cell_str(v):
            nonlocal saw_any_value
            if v is not None:
                saw_any_value = True
            return "" if v is None else str(v).replace("|", "\\|").replace("\n", " ")

        def append(line):
            nonlocal total, truncated
            if total + len(line) + 1 >= max_chars:
                buf.append(line[: max_chars - total])
                total = max_chars
                truncated = True
                return True
            buf.append(line + "\n")
            total += len(line) + 1
            return False

        if append(f"## {ws.title}"):
            return total, truncated, saw_any_value
        if append(""):
            return total, truncated, saw_any_value

        header_cells = [cell_str(v) for v in first_row]
        header_line = "| " + " | ".join(header_cells) + " |"
        divider_line = "| " + " | ".join(["---"] * len(header_cells)) + " |"
        if append(header_line):
            return total, truncated, saw_any_value
        if append(divider_line):
            return total, truncated, saw_any_value
        for row in rows_iter:
            body_cells = [cell_str(v) for v in row]
            if append("| " + " | ".join(body_cells) + " |"):
                return total, truncated, saw_any_value
        if append(""):
            return total, truncated, saw_any_value
        return total, truncated, saw_any_value

    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        sheet_count = len(wb.sheetnames)
        buf: list[str] = []
        total = 0
        truncated = False
        any_value_seen = False
        for ws in wb.worksheets:
            if truncated:
                break
            total, truncated, saw_any = _emit_sheet(ws, buf, total, truncated)
            any_value_seen = any_value_seen or saw_any
        wb.close()

        # Fallback: if every cell came back None (Python-generated xlsx with
        # no cached calc values), re-open with data_only=False to surface
        # formula text.
        if not any_value_seen and sheet_count > 0:
            buf = []
            total = 0
            truncated = False
            wb2 = openpyxl.load_workbook(path, read_only=True, data_only=False)
            for ws in wb2.worksheets:
                if truncated:
                    break
                total, truncated, _ = _emit_sheet(ws, buf, total, truncated)
            wb2.close()

        markdown = "".join(buf).rstrip()
        if not markdown.strip():
            return {"ok": False, "error": "empty_scanned"}

        return {
            "ok": True,
            "markdown": markdown,
            "pages": sheet_count,
            "chars": len(markdown),
            "truncated": truncated,
            "extractor": "openpyxl",
        }
    except Exception:
        log.exception("openpyxl extraction failed for %s", path)
        return {"ok": False, "error": "parse_error"}


def _extract_pptx(path: str, *, max_chars: int,
                  max_uncompressed_bytes: int) -> dict:
    try:
        import pptx  # noqa: F401
    except ImportError:
        return {"ok": False, "error": "not_installed"}
    return {"ok": False, "error": "not_installed"}  # replaced in Task 6
